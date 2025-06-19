# main.py
import os
from azure.cosmos import CosmosClient
from openai import OpenAI
import json
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER
import logging
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from dotenv import load_dotenv

# Initialize logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

class OpenAIHelper:
    def __init__(self):
        self.client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

    def generate_instructions(self, recipe_name: str, ingredients: list) -> Dict[str, Any]:
        ingredients_text = "\n".join([
            f"- {ing['recipe_amount'].upper()} of {ing['ingredient'].upper()}" 
            for ing in ingredients
        ])
        
        prompt = f"""
        Create detailed cooking instructions for: {recipe_name}
        
        Ingredients:
        {ingredients_text}
        
        Create a comprehensive recipe guide with:
        1. Step-by-step preparation method
        2. Cooking tips specific to this recipe
        3. Timing for each major step
        4. Key techniques required
        5. Storage and serving suggestions
        
        Return as JSON with:
        - preparation_steps: (array of strings) Detailed steps
        - cooking_tips: (array of strings) At least 3 specific tips
        - timing: (object) Time estimates for major steps
        - techniques: (array of strings) Key cooking techniques
        - storage: (string) Storage instructions
        - serving: (string) Serving suggestions
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4-turbo-preview",
                messages=[
                    {"role": "system", "content": "You are a professional chef creating detailed cooking instructions."},
                    {"role": "user", "content": prompt}
                ],
                response_format={ "type": "json_object" }
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            logger.error(f"Error generating instructions: {e}")
            raise

class PDFGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.setup_styles()

    def setup_styles(self):
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        self.section_style = ParagraphStyle(
            'Section',
            parent=self.styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.HexColor('#2E5A88')
        )

    def create_recipe_pdf(self, recipe_list: List[Dict], ai_instructions_list: List[Dict], output_path: str):
        logger.info(f"Creating PDF with {len(recipe_list)} recipes")
        
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=50
        )

        story = []
        
        for idx, (recipe, ai_instructions) in enumerate(zip(recipe_list, ai_instructions_list)):
            if idx > 0:
                story.append(PageBreak())
            
            logger.info(f"Processing recipe {idx + 1}: {recipe.get('name', 'Unknown')}")
            
            # Recipe data is directly in the recipe object
            recipe_info = recipe['data']
            
            # Title
            story.append(Paragraph(recipe['name'], self.title_style))
            story.append(Spacer(1, 20))

            # Recipe Information
            info_data = [
                ["Servings", str(recipe_info['servings'])],
                ["Total Cost", f"${recipe_info['total_cost']:.2f}"],
                ["Cost per Serving", f"${recipe_info['cost_per_serving']:.2f}"]
            ]
            info_table = Table(info_data, colWidths=[2*inch, 4*inch])
            info_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F5F5F5')),
                ('PADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(info_table)
            story.append(Spacer(1, 20))

            # Ingredients
            story.append(Paragraph("Ingredients", self.section_style))
            ingredients_data = [["INGREDIENT", "AMOUNT", "COST PER UNIT", "TOTAL COST"]]
            for ing in recipe_info['ingredients']:
                ingredients_data.append([
                    ing['ingredient'].upper(),
                    ing['recipe_amount'].upper(),
                    f"${ing['unit_cost']:.2f}",
                    f"${ing['total_cost']:.2f}"
                ])
            ing_table = Table(ingredients_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch, 1.5*inch])
            ing_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E5A88')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('PADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(ing_table)
            story.append(Spacer(1, 20))

            # AI-Generated Instructions
            if ai_instructions:
                # Preparation Steps
                story.append(Paragraph("Preparation Method", self.section_style))
                for i, step in enumerate(ai_instructions['preparation_steps'], 1):
                    story.append(Paragraph(f"{i}. {step}", self.styles['Normal']))
                story.append(Spacer(1, 20))

                # Cooking Tips
                story.append(Paragraph("Cooking Tips", self.section_style))
                for tip in ai_instructions['cooking_tips']:
                    story.append(Paragraph(f"• {tip}", self.styles['Normal']))
                story.append(Spacer(1, 20))

                # Timing Information
                story.append(Paragraph("Timing", self.section_style))
                for step, time in ai_instructions['timing'].items():
                    story.append(Paragraph(f"• {step}: {time}", self.styles['Normal']))
                story.append(Spacer(1, 20))

                # Storage and Serving
                story.append(Paragraph("Storage", self.section_style))
                story.append(Paragraph(ai_instructions['storage'], self.styles['Normal']))
                story.append(Spacer(1, 12))
                
                story.append(Paragraph("Serving Suggestions", self.section_style))
                story.append(Paragraph(ai_instructions['serving'], self.styles['Normal']))

        doc.build(story)
        logger.info(f"PDF generation completed: {output_path}")

class WordGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.size = Pt(24)
        title_font.bold = True
        
        heading_style = self.doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.size = Pt(16)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(46, 90, 136)  # #2E5A88

    def create_recipe_docx(self, recipe_list: List[Dict], ai_instructions_list: List[Dict], output_path: str):
        logger.info(f"Creating Word document with {len(recipe_list)} recipes")
        
        for idx, (recipe, ai_instructions) in enumerate(zip(recipe_list, ai_instructions_list)):
            if idx > 0:
                self.doc.add_page_break()
            
            logger.info(f"Processing recipe {idx + 1}: {recipe.get('name', 'Unknown')}")
            
            recipe_info = recipe['data']

            # Title
            title = self.doc.add_paragraph(recipe['name'], 'CustomTitle')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.doc.add_paragraph()

            # Recipe Information
            info_table = self.doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'
            
            cells = [
                ("Servings", str(recipe_info['servings'])),
                ("Total Cost", f"${recipe_info['total_cost']:.2f}"),
                ("Cost per Serving", f"${recipe_info['cost_per_serving']:.2f}")
            ]
            
            for i, (label, value) in enumerate(cells):
                row = info_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value

            self.doc.add_paragraph()

            # Ingredients
            self.doc.add_paragraph("INGREDIENTS", 'CustomHeading')
            ingredients_table = self.doc.add_table(rows=1, cols=4)
            ingredients_table.style = 'Table Grid'
            
            header_cells = ingredients_table.rows[0].cells
            headers = ["INGREDIENT", "AMOUNT", "COST PER UNIT", "TOTAL COST"]
            for i, text in enumerate(headers):
                header_cells[i].text = text
                
            for ing in recipe_info['ingredients']:
                row = ingredients_table.add_row()
                row.cells[0].text = ing['ingredient'].upper()
                row.cells[1].text = ing['recipe_amount'].upper()
                row.cells[2].text = f"${ing['unit_cost']:.2f}"
                row.cells[3].text = f"${ing['total_cost']:.2f}"

            self.doc.add_paragraph()

            if ai_instructions:
                # Preparation Steps
                self.doc.add_paragraph("Preparation Method", 'CustomHeading')
                for i, step in enumerate(ai_instructions['preparation_steps'], 1):
                    self.doc.add_paragraph(f"{i}. {step}")
                self.doc.add_paragraph()

                # Cooking Tips
                self.doc.add_paragraph("Cooking Tips", 'CustomHeading')
                for tip in ai_instructions['cooking_tips']:
                    self.doc.add_paragraph(f"• {tip}")
                self.doc.add_paragraph()

                # Timing Information
                self.doc.add_paragraph("Timing", 'CustomHeading')
                for step, time in ai_instructions['timing'].items():
                    self.doc.add_paragraph(f"• {step}: {time}")
                self.doc.add_paragraph()

                # Storage and Serving
                self.doc.add_paragraph("Storage", 'CustomHeading')
                self.doc.add_paragraph(ai_instructions['storage'])
                
                self.doc.add_paragraph("Serving Suggestions", 'CustomHeading')
                self.doc.add_paragraph(ai_instructions['serving'])

        self.doc.save(output_path)
        logger.info(f"Word document generation completed: {output_path}")